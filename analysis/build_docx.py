# -*- coding: utf-8 -*-
"""
build_docx.py — render manuscript/manuscript.md into a formatted Word document
with embedded figures and tables.  Output: manuscript/manuscript.docx
"""
import os, re, glob
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
MAN = os.path.join(ROOT, "manuscript")
FIG = os.path.join(ROOT, "results", "figures")
TBL = os.path.join(ROOT, "results", "tables")
TEAL = RGBColor(0x0E, 0x7C, 0x86)
ORANGE = RGBColor(0xDD, 0x8A, 0x33)

FIG_CAPTIONS = {
    "F1_calibration.png": "Fig 1. Cohort calibration. (A) Target vs achieved caseness prevalence. (B) Achieved total-score correlation matrix.",
    "F2_item_utility.png": "Fig 2. Per-item cross-diagnostic utility by domain. Short item codes; orange = selected anchor (λ=0.10).",
    "F3_lockbox_performance.png": "Fig 3. Panel performance. (A) Lockbox AUROC, anchor-only vs full seven-item panel. (B) Cross-diagnostic gain (panic anchor-dominant; orange = sleep).",
    "F4_ggm_network.png": "Fig 4. Partial-correlation (Gaussian graphical model) network of the seven-domain panel. Orange = sleep anchor; blue/red edges = positive/negative.",
    "F5_contribution_heatmap.png": "Fig 5. Cross-diagnostic contribution weights (standardized logistic coefficients). Off-diagonal entries indicate transdiagnostic signal.",
    "F6_value_of_sleep.png": "Fig 6. Value of the sleep domain. (A) Negligible incremental prediction of the six original domains. (B) The comorbid panel improves insomnia detection over the sleep anchor alone.",
    "F7_bootstrap_stability.png": "Fig 7. Bootstrap selection stability (300 resamples); top-three items per domain (orange = main-analysis anchor).",
    "F8_robustness.png": "Fig 8. Robustness. (A) Optimized vs random one-item-per-domain panel (lockbox AUROC). (B) Sleep anchor across 36 perturbed calibrations; 89% land on an ISI-3m item (orange) even with sleep cross-loadings removed.",
}
TABLE_CAPTIONS = {
    "T1_calibration.csv": "Table 1. Cohort calibration: target vs achieved prevalence and total-score summaries.",
    "T3_selected_panel.csv": "Table 2. Selected seven-domain panel (λ=0.10) with item utility and bootstrap selection frequency.",
    "T4_lockbox_performance.csv": "Table 3. Lockbox performance: anchor-only vs full panel AUROC/AUPRC and cross-diagnostic gain.",
    "T11_auroc_ci.csv": "Table 4. Lockbox AUROC with bootstrap 95% confidence intervals (panel vs anchor).",
    "T5_value_of_sleep.csv": "Table 5. Incremental value of adding the sleep anchor for the six original domains.",
    "T7_lambda_sensitivity.csv": "Table 6. Selected anchors across the redundancy weight λ (sensitivity).",
    "T9_bridge_centrality.csv": "Table 7. Network bridge centrality (node strength, betweenness).",
    "T12_baselines.csv": "Table 8. Baselines: random one-item panel vs optimized panel (full battery is trivially 1.0).",
    "T13_ablation.csv": "Table 9. Calibration ablation (36 conditions): selected sleep anchor and whether it is an ISI-3m item.",
    "T14_deconfounding.csv": "Table 10. Redundancy (cosine) of the selected ISI sleep anchor with PHQ-9 sleep, fatigue, mood and anhedonia items — the mechanism of the depression de-confounding.",
}

INLINE = re.compile(r"(\*\*.+?\*\*|`.+?`|\*.+?\*)")


def add_inline(p, text):
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)  # links -> text (url)
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = p.add_run(tok[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9.5)
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            r = p.add_run(tok[1:-1]); r.italic = True
        else:
            p.add_run(tok)


def shade(paragraph, hexcolor="EFF5F6"):
    pPr = paragraph._p.get_or_add_pPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor); pPr.append(sh)


def render_csv_table(doc, csv_path):
    df = pd.read_csv(csv_path)
    df = df.fillna("")
    t = doc.add_table(rows=1, cols=len(df.columns))
    t.style = "Light Grid Accent 1"
    for j, c in enumerate(df.columns):
        cell = t.rows[0].cells[j]
        cell.text = str(c)
        for run in cell.paragraphs[0].runs:
            run.bold = True; run.font.size = Pt(8)
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for j, c in enumerate(df.columns):
            cells[j].text = str(row[c])
            for run in cells[j].paragraphs[0].runs:
                run.font.size = Pt(8)
    return t


def is_table_row(line):
    return line.strip().startswith("|") and line.strip().endswith("|")


def main():
    md = open(os.path.join(MAN, "manuscript.md"), encoding="utf-8").read().splitlines()
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    # --- PLOS ONE formatting: double line spacing, continuous line numbers, page numbers ---
    from docx.enum.text import WD_LINE_SPACING
    doc.styles["Normal"].paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    _sectPr = doc.sections[0]._sectPr
    _ln = OxmlElement("w:lnNumType")
    _ln.set(qn("w:countBy"), "1"); _ln.set(qn("w:restart"), "continuous"); _ln.set(qn("w:distance"), "360")
    _sectPr.append(_ln)
    _fp = doc.sections[0].footer.paragraphs[0]; _fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run = _fp.add_run()
    _b = OxmlElement("w:fldChar"); _b.set(qn("w:fldCharType"), "begin")
    _instr = OxmlElement("w:instrText"); _instr.set(qn("xml:space"), "preserve"); _instr.text = "PAGE"
    _e = OxmlElement("w:fldChar"); _e.set(qn("w:fldCharType"), "end")
    _run._r.append(_b); _run._r.append(_instr); _run._r.append(_e)

    i = 0
    while i < len(md):
        line = md[i]
        s = line.strip()
        if s == "---":
            i += 1; continue
        # markdown table block
        if is_table_row(line):
            block = []
            while i < len(md) and is_table_row(md[i]):
                block.append(md[i]); i += 1
            rows = [[c.strip() for c in r.strip().strip("|").split("|")] for r in block]
            rows = [r for r in rows if not all(set(c) <= set("-: ") for c in r)]
            if rows:
                t = doc.add_table(rows=0, cols=len(rows[0])); t.style = "Light Grid Accent 1"
                for ri, r in enumerate(rows):
                    cells = t.add_row().cells
                    for j, c in enumerate(r[:len(cells)]):
                        cells[j].text = ""
                        add_inline(cells[j].paragraphs[0], c)
                        for run in cells[j].paragraphs[0].runs:
                            run.font.size = Pt(8)
                            if ri == 0: run.bold = True
            continue
        if s.startswith("> "):
            p = doc.add_paragraph(); shade(p)
            p.paragraph_format.left_indent = Inches(0.15)
            txt = s[2:]
            if txt.startswith("###"):
                add_inline(p, txt.lstrip("#").strip())
                for r in p.runs: r.bold = True; r.font.color.rgb = ORANGE
            else:
                add_inline(p, txt)
            i += 1; continue
        if s.startswith("# "):
            h = doc.add_heading(level=0); add_inline(h, s[2:]); i += 1; continue
        if s.startswith("## "):
            h = doc.add_heading(s[3:], level=1)
            for r in h.runs: r.font.color.rgb = TEAL
            i += 1; continue
        if s.startswith("### "):
            h = doc.add_heading(s[4:], level=2)
            for r in h.runs: r.font.color.rgb = TEAL
            i += 1; continue
        m = re.match(r"^(\d+)\.\s+(.*)", s)
        if m:
            p = doc.add_paragraph(style="List Number"); add_inline(p, m.group(2)); i += 1; continue
        if s.startswith("- "):
            p = doc.add_paragraph(style="List Bullet"); add_inline(p, s[2:]); i += 1; continue
        if s == "":
            i += 1; continue
        p = doc.add_paragraph(); add_inline(p, s); i += 1

    # ---- Tables (placed in the manuscript file per PLOS; figures are separate TIFFs) ----
    doc.add_page_break()
    h = doc.add_heading("Tables", level=1)
    for r in h.runs: r.font.color.rgb = TEAL
    for csv_name, cap in TABLE_CAPTIONS.items():
        path = os.path.join(TBL, csv_name)
        if os.path.exists(path):
            c = doc.add_paragraph(); add_inline(c, cap)
            for run in c.runs: run.bold = True; run.font.size = Pt(9.5)
            render_csv_table(doc, path)
            doc.add_paragraph("")

    # ---- Real-data validation outputs (Fig 9, Tables R1–R3) ----
    RDR = os.path.join(ROOT, "results", "realdata")
    rf = os.path.join(RDR, "RF1_realdata.png")
    if os.path.exists(rf):
        doc.add_page_break()
        h = doc.add_heading("Real-data tables (Cohort A, N = 24,292)", level=1)
        for r in h.runs: r.font.color.rgb = TEAL
        rd_tables = {
            "RT2_selected_panel.csv": "Table R1. Real-data selected three-domain panel with utility and bootstrap selection frequency.",
            "RT2b_isi_selection_freq.csv": "Table R2. Real-data ISI-item selection frequency as the sleep anchor (300 bootstraps).",
            "RT3_lockbox.csv": "Table R3. Real-data lockbox AUROC (panel vs anchor) with bootstrap 95% CIs.",
            "RT4_isi_lambda.csv": "Table R4. Real-data: selected ISI item across the redundancy weight λ.",
        }
        for csv_name, cap_txt in rd_tables.items():
            p = os.path.join(RDR, csv_name)
            if os.path.exists(p):
                c = doc.add_paragraph(); add_inline(c, cap_txt)
                for run in c.runs: run.bold = True; run.font.size = Pt(9.5)
                render_csv_table(doc, p)
                doc.add_paragraph("")

    # ---- Multi-cohort validation (Fig 10, Tables R5–R7) ----
    rf2 = os.path.join(RDR, "RF2_multicohort.png")
    if os.path.exists(rf2):
        doc.add_page_break()
        h = doc.add_heading("Multi-cohort tables (Cohorts B–D)", level=1)
        for r in h.runs: r.font.color.rgb = TEAL
        mc_tables = {
            "RT5_multicohort_summary.csv": "Table R5. Cross-cohort summary: selected sleep/anxiety items and H1 verdict across three independent open cohorts.",
            "RT6_cohortB_sri.csv": "Table R6. Cohort B (clinical, N=95): ISI-item selection frequency (200 bootstraps).",
            "RT7_cohortC_uk.csv": "Table R7. Cohort C (UK, N=1,408): five-domain panel (incl. suicidality) anchor and lockbox AUROC.",
            "RT8_cohortD.csv": "Table R8. Cohort D (BELL-001 MoA; Korean clinical insomnia, N=33): ISI-item selection — underpowered/exploratory.",
        }
        for csv_name, cap_txt in mc_tables.items():
            p = os.path.join(RDR, csv_name)
            if os.path.exists(p):
                c = doc.add_paragraph(); add_inline(c, cap_txt)
                for run in c.runs: run.bold = True; run.font.size = Pt(9.5)
                render_csv_table(doc, p)
                doc.add_paragraph("")

    # ---- Figure captions (figures uploaded SEPARATELY as TIFF per PLOS; not embedded) ----
    doc.add_page_break()
    h = doc.add_heading("Figure captions", level=1)
    for r in h.runs: r.font.color.rgb = TEAL
    _n = doc.add_paragraph()
    add_inline(_n, "Figures are provided as separate files (Fig1.tif–Fig10.tif), per PLOS ONE requirements; captions are listed here in read order.")
    for run in _n.runs: run.italic = True; run.font.size = Pt(9)
    _figcaps = [FIG_CAPTIONS[k] for k in sorted(FIG_CAPTIONS)]
    _figcaps.append("Fig 9. Real-data validation (Cohort A; Zenodo 10423537; item-level ISI, PHQ-9, GAD-7). "
                    "(A) Bootstrap selection frequency of each ISI item as the sleep anchor (orange = ISI-3m items); "
                    "(B) three-domain lockbox AUROC, anchor vs panel; (C) total-score correlations, real vs literature.")
    _figcaps.append("Fig 10. Multi-cohort validation. (A) Cohort B (SRI; US adolescents, clinical insomnia, N=95): "
                    "bootstrap selection frequency of each ISI item as the sleep anchor (orange = ISI-3m) — unstable "
                    "at this sample size. (B) Cohort C (Akram et al.; UK, N=1,408): five-domain lockbox AUROC including "
                    "the new suicidality domain (orange; insomnia measured by the SCI).")
    for _cap in _figcaps:
        _p = doc.add_paragraph(); add_inline(_p, _cap)

    out = os.path.join(MAN, "manuscript.docx")
    doc.save(out)
    print(f"[build_docx] saved {out}")


if __name__ == "__main__":
    main()
