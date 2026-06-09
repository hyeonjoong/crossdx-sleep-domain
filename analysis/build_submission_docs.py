# -*- coding: utf-8 -*-
"""
build_submission_docs.py — generate the readable Supporting-Information and
cover-letter documents for the PLOS ONE submission:
  SUBMISSION/cover_letter.docx   (from manuscript/cover_letter.md)
  SUBMISSION/S1_Protocol.docx    (from protocol/analysis_plan.md)
  SUBMISSION/S1_Checklist.docx   (adapted TRIPOD+AI reporting checklist)
"""
import os, re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
SUB = os.path.join(ROOT, "SUBMISSION")
os.makedirs(SUB, exist_ok=True)
TEAL = RGBColor(0x0E, 0x7C, 0x86)
INLINE = re.compile(r"(\*\*.+?\*\*|`.+?`)")


def add_inline(p, text):
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = p.add_run(tok[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(10)
        else:
            p.add_run(tok)


def md_to_docx(md_path, out_path, base_font=11):
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(base_font)
    for line in open(md_path, encoding="utf-8").read().splitlines():
        s = line.strip()
        if s == "---" or s == "```":
            continue
        if s.startswith("# "):
            h = doc.add_heading(s[2:], level=0)
        elif s.startswith("## "):
            h = doc.add_heading(s[3:], level=1)
            for r in h.runs: r.font.color.rgb = TEAL
        elif s.startswith("### "):
            h = doc.add_heading(s[4:], level=2)
            for r in h.runs: r.font.color.rgb = TEAL
        elif re.match(r"^\d+\.\s+", s):
            p = doc.add_paragraph(style="List Number"); add_inline(p, re.sub(r"^\d+\.\s+", "", s))
        elif s.startswith("- "):
            p = doc.add_paragraph(style="List Bullet"); add_inline(p, s[2:])
        elif s == "":
            doc.add_paragraph("")
        else:
            p = doc.add_paragraph(); add_inline(p, s)
    doc.save(out_path)
    print("  saved", os.path.basename(out_path))


# ---- TRIPOD+AI (adapted) reporting checklist ----
TRIPOD = [
    ("Title & abstract", "", "", ""),
    ("1", "Title identifies the study as developing/validating a multivariable prediction model and the target", "Yes", "Title"),
    ("2", "Abstract: structured summary of objectives, methods, results, conclusions", "Yes", "Abstract"),
    ("Introduction", "", "", ""),
    ("3a", "Background and rationale, including references to existing models", "Yes", "Introduction"),
    ("3b", "Study objectives / questions (Q1–Q3)", "Yes", "Introduction"),
    ("Methods", "", "", ""),
    ("4a", "Source of data (simulated cohort; four real open/clinical cohorts)", "Yes", "Methods: Simulated cohort; Real-data validation"),
    ("4b", "Dates/setting of data", "Yes", "Methods (cohort descriptions)"),
    ("5", "Participants / eligibility (per cohort, incl. cutoffs)", "Yes", "Methods: Real-data validation"),
    ("6", "Outcome (binary caseness per instrument cutoff) and how assessed", "Yes", "Methods"),
    ("7", "Predictors (questionnaire items) and how handled", "Yes", "Methods: optimization"),
    ("8", "Sample size", "Yes", "Methods (N per cohort)"),
    ("9", "Missing data handling", "Partial", "Methods (complete-case for real cohorts; simulation has none)"),
    ("10a", "Model: type and analytical method (one-item-per-domain optimization; L2-logistic)", "Yes", "Methods: optimization; Validation"),
    ("10b", "Model building: predictor selection (utility minus cosine-redundancy)", "Yes", "Methods"),
    ("10c", "Model output", "Yes", "Methods (caseness probability/AUROC)"),
    ("11", "Performance measures (AUROC, AUPRC, bootstrap CIs) and validation (lockbox, bootstrap)", "Yes", "Methods: Validation, uncertainty, baselines, ablation"),
    ("12", "Model fairness / sensitivity analyses (λ sensitivity; calibration ablation)", "Yes", "Methods; Results §3.3"),
    ("Open science", "", "", ""),
    ("13", "Funding", "Yes", "Funding statement"),
    ("14", "Conflicts of interest", "Yes", "Competing interests"),
    ("15", "Protocol / registration", "Yes", "S1 Protocol (pre-registered confirmatory plan)"),
    ("16", "Data availability", "Yes", "Data availability statement; S1 File"),
    ("17", "Code availability", "Yes", "S1 File; public repository"),
    ("Results", "", "", ""),
    ("18", "Participant flow / cohort descriptives", "Yes", "Results §3.1; Tables 1, R-tables"),
    ("19", "Model specification (selected items)", "Yes", "Results §3.2–3.3; Tables 2, 6"),
    ("20", "Model performance (lockbox AUROC, CIs, gains)", "Yes", "Results §3.6; Tables 3, 4"),
    ("Discussion", "", "", ""),
    ("21", "Limitations", "Yes", "Discussion: Limitations"),
    ("22", "Interpretation (emergent vs assumption-driven)", "Yes", "Discussion"),
    ("23", "Generalizability (cross-population; clinical pending)", "Yes", "Discussion; Limitations"),
    ("24", "Clinical/use implications and next steps", "Yes", "Discussion; S1 Protocol"),
]


def build_checklist():
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(10)
    h = doc.add_heading("S1 Checklist. TRIPOD+AI reporting checklist (adapted)", level=0)
    note = doc.add_paragraph()
    note.add_run("This is a development/simulation + multi-cohort validation study, not the "
                 "derivation of a deployed clinical prediction tool; items specific to clinical "
                 "deployment or to calibration of a fielded model are marked Partial/NA. Item "
                 "numbering follows the TRIPOD+AI (Collins et al., BMJ 2024) structure; consult "
                 "the official EQUATOR checklist for verbatim item wording.").italic = True
    for r in note.runs: r.font.size = Pt(9)
    doc.add_paragraph("")
    t = doc.add_table(rows=1, cols=4); t.style = "Light Grid Accent 1"
    for j, head in enumerate(["Item", "Checklist item", "Reported", "Location"]):
        c = t.rows[0].cells[j]; c.text = head
        for run in c.paragraphs[0].runs: run.bold = True; run.font.size = Pt(9)
    for item, desc, rep, loc in TRIPOD:
        cells = t.add_row().cells
        if desc == "":  # section header row
            cells[0].merge(cells[3]); cells[0].text = item
            for run in cells[0].paragraphs[0].runs: run.bold = True; run.font.color.rgb = TEAL; run.font.size = Pt(9.5)
            continue
        for j, val in enumerate([item, desc, rep, loc]):
            cells[j].text = val
            for run in cells[j].paragraphs[0].runs: run.font.size = Pt(8.5)
    doc.save(os.path.join(SUB, "S1_Checklist.docx"))
    print("  saved S1_Checklist.docx")


def main():
    print("[build_submission_docs] generating readable SI + cover letter ...")
    md_to_docx(os.path.join(ROOT, "manuscript", "cover_letter.md"), os.path.join(SUB, "cover_letter.docx"))
    md_to_docx(os.path.join(ROOT, "protocol", "analysis_plan.md"), os.path.join(SUB, "S1_Protocol.docx"))
    build_checklist()


if __name__ == "__main__":
    main()
